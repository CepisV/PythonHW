#Exel
#a
import pandas as pd

file_1 = pd.read_excel('1111.xlsx', engine='openpyxl')
file_2 = pd.read_excel('2222.xlsx', engine='openpyxl')
file_3 = pd.read_excel('3333.xlsx', engine='openpyxl')

merged_data = pd.concat([file_1, file_2, file_3], ignore_index=True)

print(merged_data)

#b
sorted_data = merged_data.sort_values(by=merged_data.columns.tolist(), ascending=False)
print(sorted_data)

#c
from openpyxl import Workbook
from openpyxl.styles import Font, Border, Side

wb = Workbook()
ws = wb.active

for r_idx, row in sorted_data.iterrows():
    for c_idx, value in enumerate(row):
        cell = ws.cell(row=r_idx+1, column=c_idx+1, value=value)

font = Font(name='Calibri', size=12, bold=True)
border = Border(left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin'))

for row in ws.iter_rows():
    for cell in row:
        cell.font = font
        cell.border = border

wb.save('sorted_data.xlsx')

#JSON
#a
import requests
import json

url = 'https://jsonplaceholder.typicode.com/todos/'
response = requests.get(url)
data = response.json()

with open('todos.json', 'w') as f:
    json.dump(data, f, indent=4)

#b
with open('todos.json', 'r') as f:
    todos = json.load(f)

print(todos)
    
#c
for idx, todo in enumerate(todos):
    with open(f'todo_{idx+1}.json', 'w') as f:
        json.dump(todo, f, indent=4)

#WORD
#a
from docx import Document

doc = Document()

doc.add_paragraph('Hello Python')

doc.save('hello_python.docx')

#b
from docx import Document

doc = Document('hello_python.docx')

for para in doc.paragraphs:
    for run in para.runs:
        if run.bold:  
            print(run.text)

#c
doc = Document()

paragraph = doc.add_paragraph()

run = paragraph.add_run('This is a new paragraph with custom font.')
run.font.name = 'Arial'
run.font.size = 12 

doc.save('formatted_text.docx')
