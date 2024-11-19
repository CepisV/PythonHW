#zad1 Excel
import openpyxl

wb = openpyxl.Workbook()
ws = wb.active

ws.append([i for i in range(1, 11)]) 
ws.append([chr(97 + i) for i in range(10)])  

wb.save("data.xlsx")
print("Файл data.xlsx создан и заполнен.")


#b
wb = openpyxl.load_workbook("data.xlsx")
ws = wb.active

row1 = [cell.value for cell in ws[1]]
row2 = [cell.value for cell in ws[2]]

print("Первая строка:", row1)
print("Вторая строка:", row2)

#c
wb_new = openpyxl.Workbook()
ws_new = wb_new.active

for i, value in enumerate(row1, start=1):
    ws_new.cell(row=i, column=1, value=value)
for i, value in enumerate(row2, start=1):
    ws_new.cell(row=i, column=2, value=value)

wb_new.save("data_vertical.xlsx")
print("Файл data_vertical.xlsx создан.")

#zad2
#a
import requests
import json

url = "https://jsonplaceholder.typicode.com/todos/1"
response = requests.get(url)
data = response.json()

with open("todo.json", "w") as file:
    json.dump(data, file, indent=4)

print("Данные сохранены в файл todo.json")

#b
with open("todo.json", "r") as file:
    todo_dict = json.load(file)

print("Содержимое JSON-файла:", todo_dict)

#c
todos = []
for i in range(1, 101):
    todos.append({"id": i, "title": f"Task {i}", "completed": False})

with open("todos.json", "w") as file:
    json.dump(todos, file, indent=4)

print("100 словарей сохранены в файл todos.json")

#zad3
#a
from docx import Document

doc = Document()
doc.add_paragraph("Hello Python")
doc.save("hello.docx")
print("Файл hello.docx создан.")

#b
doc = Document("hello.docx")
content = "\n".join([p.text for p in doc.paragraphs])

print("Содержимое файла hello.docx:")
print(content)

#c
doc_new = Document()
doc_new.add_paragraph("This is a new paragraph in a new Word file.")
doc_new.save("new_paragraph.docx")
print("Файл new_paragraph.docx создан.")
